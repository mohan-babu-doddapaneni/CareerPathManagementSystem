"""
Enhanced resume parser - supports .docx and .pdf.
Section-based extraction: skills, education, experience, certifications.
"""
import re
import io


def extract_text_from_docx(file) -> str:
    from docx import Document
    buf = file.read() if hasattr(file, 'read') else open(file, 'rb').read()
    doc = Document(io.BytesIO(buf))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return '\n'.join(parts)


def extract_text_from_pdf(file) -> str:
    buf = file.read() if hasattr(file, 'read') else open(file, 'rb').read()
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(buf))
        pages = [page.extract_text() or '' for page in reader.pages]
        text = '\n'.join(pages)
        if text.strip():
            return text
    except Exception:
        pass
    try:
        from pdfminer.high_level import extract_text as pm_extract
        return pm_extract(io.BytesIO(buf)) or ''
    except Exception:
        pass
    raise ValueError("Could not extract text from PDF.")


def extract_text(file, filename: str = '') -> str:
    ext = (filename or getattr(file, 'name', '')).lower().rsplit('.', 1)[-1]
    if hasattr(file, 'seek'):
        file.seek(0)
    if ext == 'pdf':
        return extract_text_from_pdf(file)
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(file)
    raise ValueError(f"Unsupported file type: .{ext}. Upload .docx or .pdf.")


_SECTION_RE = {
    'summary':        re.compile(r'\b(summary|profile|objective|about\s+me|overview|professional\s+summary|career\s+objective|personal\s+statement)\b', re.I),
    'experience':     re.compile(r'\b(experience|employment|work\s+history|professional\s+experience|career\s+history|internship)\b', re.I),
    'education':      re.compile(r'\b(education|academic|qualifications?|degrees?|studies|schooling|training)\b', re.I),
    'skills':         re.compile(r'\b(skills?|technical\s+skills?|core\s+competencies|competencies|expertise|technologies|tools?|proficiencies)\b', re.I),
    'certifications': re.compile(r'\b(certifications?|certificates?|credentials?|licenses?|accreditations?|professional\s+development|courses?|achievements?)\b', re.I),
    'projects':       re.compile(r'\b(projects?|portfolio|personal\s+projects?)\b', re.I),
    'awards':         re.compile(r'\b(awards?|honors?|accomplishments?|recognition|publications?)\b', re.I),
    'references':     re.compile(r'\b(references?)\b', re.I),
}


def split_sections(text: str) -> dict:
    lines = text.split('\n')
    sections = {'_header': []}
    current = '_header'
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if len(s) <= 70:
            matched = None
            for name, rx in _SECTION_RE.items():
                if rx.search(s) and len(s.split()) <= 8:
                    matched = name
                    break
            if matched:
                current = matched
                sections.setdefault(current, [])
                continue
        sections.setdefault(current, []).append(s)
    return {k: '\n'.join(v) for k, v in sections.items()}


def extract_email(text: str) -> str:
    m = re.search(r'[\w.+\-]+@[\w\-]+\.[\w.]+', text)
    return m.group(0) if m else ''


def extract_phone(text: str) -> str:
    for m in re.finditer(r'(\+?\d[\d\s\-().]{6,}\d)', text):
        digits = re.sub(r'\D', '', m.group(0))
        if 7 <= len(digits) <= 15:
            return m.group(0).strip()
    return ''


def extract_linkedin(text: str) -> str:
    m = re.search(r'linkedin\.com/in/([\w\-]+)', text, re.I)
    return 'https://linkedin.com/in/' + m.group(1) if m else ''


def extract_github(text: str) -> str:
    m = re.search(r'github\.com/([\w\-]+)', text, re.I)
    return 'https://github.com/' + m.group(1) if m else ''


_SKILLS = sorted([
    'Python', 'Java', 'JavaScript', 'TypeScript', 'C', 'C++', 'C#', 'Go', 'Golang',
    'Rust', 'Ruby', 'PHP', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl',
    'Lua', 'Haskell', 'Elixir', 'Dart', 'Julia', 'Groovy', 'Bash', 'Shell',
    'PowerShell', 'SQL', 'PL/SQL', 'T-SQL', 'COBOL', 'Assembly', 'VB.NET',
    'HTML', 'CSS', 'SASS', 'SCSS', 'LESS', 'React', 'Angular', 'Vue.js', 'Vue',
    'Next.js', 'Nuxt.js', 'Svelte', 'jQuery', 'Bootstrap', 'Tailwind CSS', 'TailwindCSS',
    'Material UI', 'Chakra UI', 'Webpack', 'Vite', 'Babel', 'Rollup',
    'Redux', 'MobX', 'Zustand', 'GraphQL', 'REST API', 'RESTful', 'WebSocket',
    'gRPC', 'OAuth', 'JWT', 'OpenAPI', 'Swagger', 'SOAP',
    'Django', 'Flask', 'FastAPI', 'Node.js', 'Express.js', 'Express', 'NestJS',
    'Spring Boot', 'Spring', 'Hibernate', 'Laravel', 'Symfony', 'Rails',
    'Ruby on Rails', 'ASP.NET', 'ASP.NET Core', '.NET', '.NET Core', 'MVC',
    'Celery', 'Gunicorn', 'Nginx', 'Apache',
    'MySQL', 'PostgreSQL', 'SQLite', 'MongoDB', 'Redis', 'Elasticsearch',
    'Cassandra', 'DynamoDB', 'Firebase', 'Supabase', 'Oracle', 'SQL Server',
    'MariaDB', 'Neo4j', 'InfluxDB', 'Couchbase', 'HBase', 'Memcached',
    'RabbitMQ', 'Apache Kafka', 'Kafka',
    'AWS', 'Azure', 'Google Cloud', 'GCP', 'Heroku', 'Railway', 'Vercel',
    'Netlify', 'DigitalOcean', 'Docker', 'Kubernetes', 'Terraform', 'Ansible',
    'Chef', 'Puppet', 'Jenkins', 'GitHub Actions', 'GitLab CI', 'CircleCI',
    'Travis CI', 'ArgoCD', 'Helm', 'Prometheus', 'Grafana', 'Datadog',
    'New Relic', 'CloudFormation', 'Pulumi',
    'Machine Learning', 'Deep Learning', 'Neural Networks', 'NLP',
    'Natural Language Processing', 'Computer Vision', 'TensorFlow', 'PyTorch',
    'Keras', 'scikit-learn', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn',
    'Plotly', 'SciPy', 'OpenCV', 'NLTK', 'spaCy', 'Hugging Face', 'LangChain',
    'XGBoost', 'LightGBM', 'Random Forest', 'Data Science', 'Data Analysis',
    'Data Engineering', 'Big Data', 'Apache Spark', 'Spark', 'Hadoop', 'Airflow',
    'dbt', 'Snowflake', 'Power BI', 'Tableau', 'Looker', 'Excel',
    'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN', 'Jira', 'Confluence',
    'Trello', 'Asana', 'Notion', 'Slack', 'VS Code', 'IntelliJ', 'Eclipse',
    'PyCharm', 'Postman', 'Insomnia', 'Figma', 'Sketch', 'Adobe XD',
    'Linux', 'Unix', 'Windows Server', 'macOS', 'Vim',
    'Unit Testing', 'Integration Testing', 'TDD', 'BDD', 'Selenium', 'Cypress',
    'Playwright', 'Jest', 'Mocha', 'Chai', 'Pytest', 'JUnit', 'TestNG',
    'Performance Testing', 'SoapUI',
    'Agile', 'Scrum', 'Kanban', 'SAFe', 'Waterfall', 'DevOps', 'CI/CD',
    'SDLC', 'Design Patterns', 'Microservices', 'Serverless', 'Event-Driven',
    'SOLID', 'Clean Code', 'Domain-Driven Design',
    'Leadership', 'Communication', 'Teamwork', 'Problem Solving',
    'Critical Thinking', 'Project Management', 'Time Management',
    'Collaboration', 'Adaptability', 'Creativity', 'Analytical Thinking',
    'Attention to Detail', 'Mentoring', 'Presentation', 'Public Speaking',
    'Cybersecurity', 'Network Security', 'Penetration Testing', 'OWASP',
    'Encryption', 'IAM', 'SSO', 'LDAP', 'Active Directory',
    'iOS', 'Android', 'React Native', 'Flutter', 'Xamarin', 'Ionic',
    'Objective-C', 'Xcode', 'Android Studio',
    'Blockchain', 'Ethereum', 'Solidity', 'Web3', 'Smart Contracts',
    'Arduino', 'Raspberry Pi', 'IoT', 'Embedded Systems',
    'Unity', 'Unreal Engine', 'Game Development', 'Three.js',
    'Salesforce', 'SAP', 'ServiceNow', 'WordPress', 'Shopify',
], key=lambda x: -len(x))


def extract_skills(text: str, skills_section: str = '') -> list:
    combined = skills_section + '\n' + text
    found = set()
    for skill in _SKILLS:
        pattern = r'(?<![A-Za-z0-9._\-])' + re.escape(skill) + r'(?![A-Za-z0-9._\-])'
        if re.search(pattern, combined, re.IGNORECASE):
            found.add(skill)
    return sorted(found)


_DEGREE_RES = [
    (re.compile(r'\bPh\.?D\.?\b', re.I), 'PhD'),
    (re.compile(r'\bDoctor(?:ate)?\s+(?:of|in)\s+([\w\s&]+)', re.I), 'Doctorate'),
    (re.compile(r'\bM\.?B\.?A\.?\b', re.I), 'MBA'),
    (re.compile(r'\bM\.?S(?:c)?\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Master of Science'),
    (re.compile(r'\bM\.?Eng\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Master of Engineering'),
    (re.compile(r'\bMaster(?:s)?\s+(?:of|in)\s+([\w\s&]{2,50})', re.I), 'Masters'),
    (re.compile(r'\bM\.?E\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Master of Engineering'),
    (re.compile(r'\bB\.?Tech\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Bachelor of Technology'),
    (re.compile(r'\bB\.?E\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Bachelor of Engineering'),
    (re.compile(r'\bB\.?S(?:c)?\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Bachelor of Science'),
    (re.compile(r'\bB\.?Eng\.?\b(?:\s+(?:in|of)\s+([\w\s&]{2,50}))?', re.I), 'Bachelor of Engineering'),
    (re.compile(r'\bBachelor(?:s)?\s+(?:of|in)\s+([\w\s&]{2,50})', re.I), 'Bachelors'),
    (re.compile(r'\bAssociate(?:s)?\s+(?:of|in)\s+([\w\s&]{2,40})', re.I), 'Associate'),
    (re.compile(r'\bDiploma\s+(?:in|of)\s+([\w\s&]{2,50})', re.I), 'Diploma'),
    (re.compile(r'\bHigh\s+School\b|\bSecondary\s+School\b', re.I), 'High School'),
    (re.compile(r'\bMatriculation\b|\bIntermediate\b', re.I), 'Matriculation'),
]

# Labels where a standalone match (no field of study captured) is reliable
_STRONG_DEGREE_LABELS = {
    'PhD', 'Doctorate', 'MBA', 'Master of Science', 'Master of Engineering',
    'Masters', 'Bachelor of Technology', 'Bachelor of Engineering',
    'Bachelor of Science', 'Bachelors', 'High School', 'Matriculation',
    'Associate', 'Diploma',
}

_INSTITUTION_KWS = [
    'university', 'college', 'institute', 'school', 'academy', 'polytechnic',
    'institution', 'faculty', 'campus', 'iit', 'iiit', 'nit', 'mit', 'stanford',
]

_YEAR_RE = re.compile(r'\b(19|20)\d{2}\b')
_DATE_RANGE_RE = re.compile(
    r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\s+\d{4}|\d{4})\s*[-\u2013\u2014to]+\s*'
    r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\s+\d{4}|\d{4}|Present|Current|Now)',
    re.I
)


_INLINE_DATE_RE = re.compile(
    r'\(?\s*(\d{4})\s*[-\u2013\u2014\u2015~to]+\s*(\d{4}|Present|Current|Now)\s*\)?',
    re.I
)


def _parse_edu_line(line: str):
    """
    Parse a single-line education entry like:
      'B.Tech in Computer Science (2016-2020), Jawaharlal Nehru Technological University, India.'
      'Masters of Science in Computer Science (2023-2025), Clark University, Worcester.'
    Returns (degree_label, field, institution, year) or (None, ...) if no degree found.
    """
    degree_label = ''
    field = ''
    match_end = 0

    for rx, label in _DEGREE_RES:
        m = rx.search(line)
        if m:
            degree_label = label
            match_end = m.end()
            try:
                cap = m.group(1).strip() if m.lastindex else ''
                # Stop captured field at commas, parens, or year
                cap = re.split(r'[,(]|\b(19|20)\d{2}\b', cap)[0].strip()
                # If "X in Y" is captured, take Y (the actual subject)
                if re.search(r'\bin\b', cap, re.I):
                    cap = re.split(r'\bin\b', cap, flags=re.I)[-1].strip()
                cap_clean = re.sub(r'\b(and|the|of|in|at|for|with|degree)\b', '', cap, flags=re.I).strip()
                if cap_clean and len(cap_clean.split()) <= 5 and not _YEAR_RE.search(cap_clean) and len(cap_clean) >= 3:
                    field = cap_clean.title()
            except Exception:
                pass
            break

    if not degree_label:
        return None, '', '', ''

    # Extract year from current line
    year = ''
    dm = _INLINE_DATE_RE.search(line)
    if dm:
        y1, y2 = dm.group(1), dm.group(2)
        year = f'{y1} - {y2}' if y2.isdigit() else f'{y1} - {y2.title()}'

    # Extract institution: part of line AFTER date range (or after degree match if no date)
    institution = ''
    remainder = line
    if dm:
        remainder = line[dm.end():]
    elif match_end:
        remainder = line[match_end:]

    # Clean up remainder: strip punctuation, "Degree" noise word, leading prepositions
    remainder = re.sub(r'\bDegree\b', '', remainder, flags=re.I)
    remainder = re.sub(r'^[\s,;:|\-.()\[\]]+|[\s,;:|\-.()\[\]]+$', '', remainder).strip()
    remainder = re.sub(r'^(in|of|at|from)\s+', '', remainder, flags=re.I).strip()
    # Strip trailing year or country abbreviation like ", 2022" or ", India"
    remainder = re.sub(r',?\s*\b(19|20)\d{2}\b\s*$', '', remainder).strip()
    remainder = re.sub(r'^[\s,;:|\-]+|[\s,;:|\-]+$', '', remainder).strip()

    # If remainder has institution keywords, use it
    if remainder and len(remainder) > 4:
        has_kw = any(kw in remainder.lower() for kw in _INSTITUTION_KWS)
        words = [w for w in remainder.split() if w]
        title_ratio = sum(1 for w in words if w[0].isupper()) / max(len(words), 1)
        # Accept as institution only if: has an institution keyword, OR 3+ words mostly title-case
        if has_kw or (len(words) >= 3 and title_ratio >= 0.6):
            institution = remainder

    return degree_label, field, institution, year


def extract_education(text: str, edu_section: str = '') -> list:
    search = edu_section if edu_section.strip() else text
    lines = [l.strip() for l in search.split('\n') if l.strip()]
    results = []
    seen_labels = set()

    for i, line in enumerate(lines):
        degree_label, field, institution, year = _parse_edu_line(line)
        if not degree_label:
            continue

        # If single-line parse didn't get institution/year, scan nearby lines
        if not institution or not year:
            ctx_lines = lines[max(0, i + 1): min(len(lines), i + 4)]
            for cl in ctx_lines:
                if not year:
                    dm = _INLINE_DATE_RE.search(cl)
                    if not dm:
                        yrs = _YEAR_RE.findall(cl)
                        if yrs:
                            year = ' - '.join(sorted(set(yrs[:2])))
                    else:
                        y1, y2 = dm.group(1), dm.group(2)
                        year = f'{y1} - {y2}' if y2.isdigit() else f'{y1} - {y2.title()}'
                if not institution:
                    for kw in _INSTITUTION_KWS:
                        if kw in cl.lower() and len(cl) > len(kw) + 2:
                            institution = re.sub(r'^[\s,;:|\-]+|[\s,;:|\-]+$', '', cl).strip()
                            break

        degree_str = degree_label + (' in ' + field if field else '')

        # Dedup: upgrade existing entry if this has more info, otherwise skip
        duplicate = False
        for existing in results:
            base_existing = existing['degree'].split(' in ')[0].strip()
            if base_existing == degree_label:
                duplicate = True
                if field and ' in ' not in existing['degree']:
                    existing['degree'] = degree_str
                if institution and not existing['institution']:
                    existing['institution'] = institution
                if year and not existing['year']:
                    existing['year'] = year
                break

        if not duplicate:
            seen_labels.add(degree_label)
            results.append({'degree': degree_str, 'institution': institution, 'year': year})

    return results


_TITLE_RES = [
    re.compile(
        r'\b(?:Senior|Sr\.?|Junior|Jr\.?|Lead|Principal|Staff|Associate|'
        r'Head\s+of|VP\s+of|Director\s+of|Chief|CTO|CEO|COO|CFO|CIO|CISO)?\s*'
        r'(?:Software|Full[\-\s]?Stack|Frontend|Back[\-\s]?end|DevOps|Cloud|Mobile|'
        r'Data|ML|AI|Research|Product|Project|Program|Business|Systems?|Network|'
        r'Security|QA|Test|Platform|Infrastructure|Embedded|Web|Game|Database|'
        r'Machine\s+Learning|Cyber)\s*'
        r'(?:Engineer(?:ing)?|Developer|Architect|Scientist|Analyst|Manager|'
        r'Lead|Director|Specialist|Consultant|Designer|Administrator|Intern|Coordinator)\b',
        re.I
    ),
    re.compile(
        r'\b(?:Full[\-\s]?Stack Developer|Software Engineer|Data Scientist|'
        r'DevOps Engineer|Cloud Architect|Product Manager|Scrum Master|Tech Lead|'
        r'Engineering Manager|Site Reliability Engineer|SRE|Platform Engineer|'
        r'Solutions Architect|IT Manager|Systems Administrator|Network Engineer|'
        r'Security Analyst|Business Analyst|Data Analyst|UI/UX Designer|'
        r'Frontend Developer|Backend Developer|Web Developer)\b',
        re.I
    ),
]


def _calc_years(start: str, end: str) -> float:
    from datetime import date
    sy = _YEAR_RE.findall(start)
    ey = _YEAR_RE.findall(end)
    if not sy:
        return 0.0
    s = int(sy[0])
    if ey:
        e = int(ey[0])
    elif re.search(r'present|current|now', end, re.I):
        e = date.today().year
    else:
        return 0.0
    return max(0.0, float(e - s))


def extract_experience(text: str, exp_section: str = '') -> list:
    search = exp_section if exp_section.strip() else text
    lines = [l.strip() for l in search.split('\n') if l.strip()]
    results = []
    for i, line in enumerate(lines):
        title = ''
        for rx in _TITLE_RES:
            m = rx.search(line)
            if m:
                title = m.group(0).strip()
                break
        if not title:
            continue
        ctx_lines = lines[i: min(len(lines), i + 5)]
        ctx = '\n'.join(ctx_lines)
        date_range = ''
        years = 0.0
        dm = _DATE_RANGE_RE.search(ctx)
        if dm:
            date_range = dm.group(0)
            years = _calc_years(dm.group(1), dm.group(2))
        company = re.sub(re.escape(title), '', line, flags=re.I).strip()
        company = re.sub(r'^[\s,|@\-at:]+|[\s,|@\-:]+$', '', company).strip()
        if not company and i + 1 < len(lines):
            nxt = lines[i + 1]
            if not _DATE_RANGE_RE.search(nxt) and not any(rx.search(nxt) for rx in _TITLE_RES) and len(nxt) < 80:
                company = nxt
        entry = {'title': title, 'company': company, 'date_range': date_range, 'years': years}
        if not any(e['title'] == entry['title'] and e['company'] == entry['company'] for e in results):
            results.append(entry)
    return results[:10]


_CERT_LINE_RES = [
    re.compile(r'\b(AWS|Azure|Google\s+Cloud|GCP)\s+(Certified|Professional|Associate|Practitioner|Developer|Architect|Administrator|DevOps)', re.I),
    re.compile(r'\b(CISSP|CISM|CISA|CEH|OSCP|CompTIA|Security\+|Network\+|Cloud\+|CySA\+|CASP\+)', re.I),
    re.compile(r'\b(PMP|CAPM|PMI|PRINCE2|ITIL|TOGAF)\b', re.I),
    re.compile(r'\b(CSM|PSM|SAFe)\s*(Certified|Practitioner|Professional|Scrum)?', re.I),
    re.compile(r'\b(CKA|CKAD|CKS)\b', re.I),
    re.compile(r'\bDocker\s+Certified\b', re.I),
    re.compile(r'\b(HashiCorp|Terraform)\s+(Associate|Certified|Professional)\b', re.I),
    re.compile(r'\b(Oracle|Java)\s+Certified\b', re.I),
    re.compile(r'\b(Red\s+Hat|RHCE|RHCSA)\b', re.I),
    re.compile(r'\bSalesforce\s+(Certified|Administrator|Developer|Architect)\b', re.I),
    re.compile(r'\b(TensorFlow\s+Developer|Databricks\s+Certified|Cloudera\s+Certified)\b', re.I),
    re.compile(r'\bMicrosoft\s+(Certified|MCP|MCSE|MCSA|MVP)\b', re.I),
    re.compile(r'\bCertified\s+[\w\s]{3,50}(?:Professional|Engineer|Developer|Analyst|Administrator|Practitioner|Specialist|Associate|Expert)\b', re.I),
    re.compile(r'\b(Professional\s+Certificate|Nanodegree)\b', re.I),
    re.compile(r'\b(Coursera|edX|Udacity|LinkedIn\s+Learning|Udemy)\s*[-:]\s*[\w\s]{5,60}', re.I),
]

_GENERIC_CERT_RE = re.compile(r'\b(certificate|certification|certified|license|accreditation|credential)\b', re.I)


def extract_certifications(text: str, cert_section: str = '') -> list:
    results = []
    seen = set()

    def add_cert(name, year='', issuer=''):
        key = re.sub(r'\s+', ' ', name.lower().strip())
        if key in seen or len(key) < 5:
            return
        seen.add(key)
        results.append({'name': name.strip(), 'year': year, 'issuer': issuer})

    for section_text, is_cert in [(cert_section, True), (text, False)]:
        if not section_text:
            continue
        for line in section_text.split('\n'):
            line = line.strip()
            if not line or len(line) < 5:
                continue
            clean = re.sub(r'^[-*>\u2022\u25cf\u25aa]\s*', '', line).strip()
            matched = False
            for rx in _CERT_LINE_RES:
                if rx.search(clean):
                    year_m = _YEAR_RE.search(clean)
                    year = year_m.group(0) if year_m else ''
                    name = _YEAR_RE.sub('', clean).strip().rstrip('.,;')
                    name = re.sub(r'\s+', ' ', name).strip()
                    add_cert(name, year)
                    matched = True
                    break
            if not matched and is_cert and _GENERIC_CERT_RE.search(clean) and len(clean) < 150:
                year_m = _YEAR_RE.search(clean)
                year = year_m.group(0) if year_m else ''
                name = _YEAR_RE.sub('', clean).strip().rstrip('.,;')
                name = re.sub(r'\s+', ' ', name).strip()
                add_cert(name, year)
        if is_cert:
            break
    return results


def extract_summary(text: str, summary_section: str = '') -> str:
    if summary_section.strip():
        return summary_section.strip()[:1000]
    for para in text.split('\n\n'):
        para = para.strip()
        if len(para) > 80 and len(para.split()) > 10:
            if not any(rx.search(para) for rx in _SECTION_RE.values()):
                return para[:500]
    return ''


def parse_resume(file, filename: str = '') -> dict:
    if not filename:
        filename = getattr(file, 'name', 'resume.docx')
    raw_text = extract_text(file, filename)
    secs = split_sections(raw_text)
    return {
        'raw_text':       raw_text,
        'email':          extract_email(raw_text),
        'phone':          extract_phone(raw_text),
        'linkedin':       extract_linkedin(raw_text),
        'github':         extract_github(raw_text),
        'summary':        extract_summary(raw_text, secs.get('summary', '')),
        'skills':         extract_skills(raw_text, secs.get('skills', '')),
        'education':      extract_education(raw_text, secs.get('education', '')),
        'experience':     extract_experience(raw_text, secs.get('experience', '')),
        'certifications': extract_certifications(raw_text, secs.get('certifications', '')),
    }
